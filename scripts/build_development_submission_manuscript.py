#!/usr/bin/env python3
"""Build a Development-targeted manuscript source and DOCX.

This script converts Markdown/LaTeX source into editable Word equations,
removes internal draft-only sections, reframes the narrative around
developmental convergence, and replaces PMID-heavy references with a
Development-like author-year reference list generated from PubMed metadata.
"""

from __future__ import annotations

import re
import subprocess
import textwrap
from pathlib import Path
from typing import Iterable
import xml.etree.ElementTree as ET
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import format_references_from_pubmed as refutil


ROOT = Path(__file__).resolve().parents[2]
MANUSCRIPT = ROOT / "Project" / "manuscript" / "granule_cell_convergence_manuscript_draft.md"
OUT_MD = ROOT / "Project" / "manuscript" / "granule_cell_convergence_Development_submission_draft.md"
OUT_DOCX = ROOT / "Project" / "manuscript" / "granule_cell_convergence_Development_submission_draft.docx"
REF_DIR = ROOT / "Project" / "manuscript" / "References"
METHODS_CLEAN = ROOT / "Project" / "manuscript" / "development_methods_clean.md"
REFERENCE_DOCX = ROOT / "Project" / "manuscript" / "development_submission_reference.docx"


TITLE = "Distinct Dentate and Cerebellar Granule-Cell Lineages Converge on Granule-like Assembly through Niche and Circuit Constraints"
SHORT_TITLE = "Integrated granule-cell convergence"
SUMMARY_STATEMENT = (
    "Distinct dentate and cerebellar granule-cell lineages preserve regional identity "
    "while converging on a granule-like assembly state shaped by niche maturation "
    "signals and sparse-expansion circuit constraints."
)

ABSTRACT = """Granule cells in the cerebellum and dentate gyrus share compact excitatory input-expansion morphologies despite arising from different developmental territories and maturing in different circuits. Similar morphology could reflect hidden shared fate, convergent postmitotic construction, niche signaling, circuit constraint, or their integration. We tested these alternatives by combining a strict 10-dataset primary core spanning mouse dentate, mouse/human cerebellar, and human dentate/hippocampal resources with ortholog-aware rank-meta modeling, named-comparator tests, pseudotime/stage analysis, ligand-receptor prediction, regulatory-compatibility scoring, morphology/activity validation, and sparse expansion-coding simulation. Dentate and cerebellar granule-cell candidates remained regionally identity-separated, arguing against one universal granule-cell fate program. Instead, downstream neurite/morphology and synaptic/excitability modules converged more strongly than upstream fate/niche modules, and an identity-coupled configuration score was broadly positive across primary-core contrasts. A conservative seed set comprising GPM6A, NFIB, NFIA, KCNK1, RFX3 and GABRA2 marked a reusable assembly and maturation toolkit rather than an exclusive granule-cell barcode. TGF-beta/BDNF/SMAD/MAPK programs behaved as stage-windowed maturation/readiness overlays, while sender-receiver analysis nominated branch-specific developmental niche cues including cerebellar Purkinje-to-granule SHH/IGF1 and dentate SGZ astrocyte/vascular/immune interactions. NeuroMorpho, DANDI and sparse expansion-coding analyses supported compact input sampling under activity and resource constraints. An evidence-weighted hypothesis matrix disfavored a hidden shared-fate explanation and supported an integrated convergence model combining identity-coupled assembly, niche maturation signals and circuit constraints. These analyses suggest that compact granule-cell morphology emerges when distinct regional lineages deploy related postmitotic assembly machinery within compatible developmental windows and sparse-expansion circuit architectures."""

KEY_WORDS = (
    "granule cell; dentate gyrus; cerebellum; developmental convergence; "
    "postmitotic assembly; neuronal morphology; single-cell transcriptomics; "
    "TGF-beta; BDNF; sparse coding; pattern separation"
)

INTRODUCTION = """Granule cells are among the most abundant neurons in the mammalian brain, but "granule cell" is not a single lineage identity. Cerebellar granule cells arise from hindbrain/rhombic-lip progenitors, expand in the external granule layer, and mature under cerebellar niche signals including SHH. Dentate gyrus granule cells arise in a telencephalic hippocampal neurogenic context and continue to pass through immature-to-mature states during postnatal and, in rodents, adult neurogenesis. Despite these distinct developmental histories, both populations adopt compact excitatory-neuron designs with small somata, dense packing, restricted input-sampling structures, and roles in transforming input spaces into sparse downstream representations (Leutgeb et al., 2007; Bird et al., 2024).

This resemblance poses a developmental mechanism question. One possibility is that cerebellar and dentate granule cells share a hidden molecular fate. A second is that distinct lineage programs independently recruit overlapping downstream machinery for neurite outgrowth, synaptogenesis, excitability tuning, adhesion/guidance and maturation. A third is that local niche signals and circuit-level input-expansion constraints favor similar compact architectures even when lineage identity remains different. These possibilities are not mutually exclusive, but they make different predictions: a shared-fate model should reveal broad granule-cell specificity, whereas a developmental-convergence model should preserve regional identity while revealing later reuse of assembly modules.

This question is sharpened by earlier work showing that cerebellar conditioned medium suppresses hippocampal granule-cell proliferation and promotes differentiation, with TGF-beta2, BDNF, SMAD and MAPK signaling implicated as candidate mediators (Lu et al., 2005). That study suggested that extrinsic cerebellar factors can push hippocampal granule-lineage cells toward differentiation. Here, we revisit the same biological question at genome-scale resolution, asking whether distinct granule-cell lineages enter related developmental assembly states.

We therefore integrated a strict 10-dataset primary transcriptomic core with ortholog-aware rank-meta modeling, pseudotime and stage analysis, named non-granule comparators, niche/pathway scoring, focused sender-receiver ligand-receptor prediction, provisional regulatory-compatibility scoring, external morphology and activity validation, hierarchical evidence synthesis, and sparse expansion-coding simulation. Cerebellar and dentate granule cells remain regionally identity-separated, but they partially converge on downstream postmitotic construction modules. Similar granule-cell morphology is therefore modeled as the output of identity-coupled assembly configuration under developmental timing, niche and circuit constraints, rather than as a single shared fate identity or a simple morphology gene list."""


RESULT_REPLACEMENTS = {
    "### Result 1. A strict 10-dataset primary core frames the granule-cell convergence problem": "### A strict 10-dataset primary core frames a developmental convergence test",
    "### Result 2. Ortholog-aware rank-meta modeling identifies a conservative shared candidate set": "### Ortholog-aware rank-meta modeling identifies a conservative shared assembly-candidate set",
    "### Result 3. Shared signals are constrained by named-comparator specificity": "### Named comparators separate shared assembly modules from granule-cell identity",
    "### Result 4. Candidate granule populations show an identity-coupled transcriptomic assembly configuration": "### An identity-coupled transcriptomic configuration captures postmitotic assembly",
    "### Result 5. Stage-windowed niche signaling and sparse-coding constraints refine the convergence model": "### Stage-windowed niche signaling and sparse-coding constraints refine the 2005 model",
    "### Result 6. Focused sender-receiver ligand-receptor prediction nominates testable niche cues": "### Sender-receiver ligand-receptor prediction nominates testable developmental niche cues",
    "### Result 7. Evidence-weighted hypothesis comparison favors an integrated convergence model": "### Evidence-weighted hypothesis comparison favors an integrated convergence model",
}

DISCUSSION_REPLACEMENTS = {
    "### A convergent assembly configuration, not a shared fate identity": "### Developmental convergence rather than shared fate identity",
    "### How sequencing supports and revises the 2005 conditioned-medium model": "### The 2005 conditioned-medium result as a stage-window clue",
    "### Morphology is partly transcriptomic, but not fully encoded as geometry": "### A transcriptomic assembly state, not a morphology gene list",
    "### Why compact sparse expansion is a useful design principle": "### Circuit constraints explain why compact granule-like designs can emerge",
    "### Experimental predictions": "### Experimental tests",
}


def section_span(text: str, start_heading: str, next_heading: str | None = None) -> tuple[int, int]:
    start = text.index(start_heading)
    if next_heading is None:
        return start, len(text)
    end = text.index(next_heading, start + len(start_heading))
    return start, end


def replace_section(text: str, start_heading: str, next_heading: str, replacement: str) -> str:
    start, end = section_span(text, start_heading, next_heading)
    return text[:start] + replacement.rstrip() + "\n\n" + text[end:]


def strip_internal_sections(text: str) -> str:
    for start, nxt in [
        ("## Figure Backbone", "## Introduction"),
        ("## Claim Language Checklist", "## Draft To-Do"),
        ("## Draft To-Do", None),
    ]:
        if start not in text:
            continue
        s = text.index(start)
        e = text.index(nxt, s) if nxt and nxt in text[s:] else len(text)
        text = text[:s] + text[e:]
    return text


def strip_pmids(text: str) -> str:
    text = re.sub(r";\s*PMID:\s*\d+", "", text)
    text = re.sub(r"\s*PMID:\s*\d+\.?", "", text)
    return text


def normalize_gene_backticks(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def normalize_table_refs(text: str) -> str:
    text = re.sub(r"\bTable_S0*([0-9]+)\b", r"Table S\1", text)
    text = re.sub(r"\bTables?_S0*([0-9]+)\b", r"Table S\1", text)
    text = re.sub(r"\bTable S([0-9]+)\s+through\s+Table S([0-9]+)", r"Tables S\1-S\2", text)
    text = re.sub(r"\bTable S([0-9]+)\s+to\s+Table S([0-9]+)", r"Tables S\1-S\2", text)
    return text


def to_page_range(pages: str) -> str:
    pages = pages.replace("-", "-")
    match = re.match(r"^(\d+)-(\d+)$", pages)
    if not match:
        return pages
    start, end = match.groups()
    if len(end) < len(start):
        end = start[: len(start) - len(end)] + end
    return f"{start}-{end}"


def development_author(author: dict[str, str]) -> str:
    if "collective" in author:
        return author["collective"]
    initials = refutil.initials_with_periods(author.get("initials", ""))
    return refutil.clean_text(f"{author.get('last', '')}, {initials}")


def join_dev_authors(authors: list[str]) -> str:
    if not authors:
        return ""
    if len(authors) == 1:
        return authors[0]
    if len(authors) == 2:
        return f"{authors[0]} and {authors[1]}"
    if len(authors) > 10:
        return ", ".join(authors[:10]) + ", et al."
    return ", ".join(authors[:-1]) + " and " + authors[-1]


def format_development_reference(record: dict[str, object]) -> str:
    authors = join_dev_authors([development_author(a) for a in record["authors"]])  # type: ignore[index]
    journal = str(record["journal_iso"] or record["journal"])
    volume = str(record["volume"])
    issue = f"({record['issue']})" if record["issue"] else ""
    pages = to_page_range(str(record["pages"]))
    doi = f" doi:{record['doi']}." if record["doi"] else ""
    return f"{authors} ({record['year']}). {record['title']}. {journal} {volume}{issue}, {pages}.{doi}"


def build_development_references(markdown: str) -> str:
    pmids = refutil.extract_pmids(markdown)
    xml_text = (REF_DIR / "pubmed_efetch.xml").read_text(encoding="utf-8")
    records = refutil.parse_pubmed(xml_text)
    ordered = [records[pmid] for pmid in pmids if pmid in records]
    alpha = sorted(
        ordered,
        key=lambda rec: (
            str(rec["authors"][0].get("last", rec["authors"][0].get("collective", ""))) if rec["authors"] else "",
            str(rec["year"]),
            str(rec["title"]),
        ),
    )
    refs = ["## References", ""]
    refs.extend(format_development_reference(rec) for rec in alpha)
    return "\n\n".join(refs) + "\n\n"


def build_submission_markdown() -> str:
    original = MANUSCRIPT.read_text(encoding="utf-8")
    front = f"""# {TITLE}

## Short Title

{SHORT_TITLE}

## Summary Statement

{SUMMARY_STATEMENT}

## Abstract

{ABSTRACT}

## Key words

{KEY_WORDS}
"""
    intro = f"## Introduction\n\n{INTRODUCTION}\n\n"

    body_start = original.index("## Results")
    body = original[body_start:]
    body = strip_internal_sections(body)
    body = replace_section(body, "## References With PMID", "## Figure Legends", build_development_references(original))

    for old, new in RESULT_REPLACEMENTS.items():
        body = body.replace(old, new)
    for old, new in DISCUSSION_REPLACEMENTS.items():
        body = body.replace(old, new)

    body = body.replace(
        "Cerebellar and dentate granule cells present a paradox: they are not the same cell type, but they share a compact excitatory input-expansion design.",
        "The developmental-convergence model predicts that cerebellar and dentate granule cells should preserve branch-specific origin while sharing later assembly logic. These cells present a useful test case because they are not the same cell type, but they share a compact excitatory input-expansion design.",
    )
    body = body.replace(
        "We next asked whether the primary core contained a reproducible shared molecular layer after species, platform, and dataset effects were reduced.",
        "Having defined a primary-core comparison, we next asked whether the data contained a reproducible shared molecular layer after species, platform and dataset effects were reduced, and whether that layer was more consistent with fate identity or postmitotic assembly.",
    )
    body = body.replace(
        "This filtering produced a compact manuscript-facing candidate structure rather than a simple long gene list:",
        "This filtering produced a compact developmental-mechanism candidate structure rather than a simple long gene list:",
    )
    body = body.replace(
        "To formalize the convergence model, we defined a transcriptomic configuration score combining two components: construction-over-niche balance and regional fate polarity",
        "To formalize convergence as a developmental state rather than a marker list, we defined a transcriptomic configuration score combining two components: construction-over-niche balance and regional fate polarity",
    )
    body = body.replace(
        "The 2005 conditioned-medium result motivated a pathway-readiness analysis of TGF-beta/BDNF/SMAD/MAPK and related niche pathways",
        "The 2005 conditioned-medium result motivated a developmental-window analysis of TGF-beta/BDNF/SMAD/MAPK and related niche pathways",
    )
    body = body.replace(
        "To move beyond pathway-readiness scores, we built a focused sender-receiver ligand-receptor analysis for the two primary datasets with suitable niche and receiver annotations.",
        "To move from pathway readiness toward experimentally testable developmental niche hypotheses, we built a focused sender-receiver ligand-receptor analysis for the two primary datasets with suitable niche and receiver annotations.",
    )
    body = body.replace(
        "Source-paper PMIDs embedded in NeuroMorpho records were retained during curation where available;",
        "Source-paper identifiers embedded in NeuroMorpho records were retained during curation where available;",
    )
    body = body.replace(
        "The central conclusion is that dentate and cerebellar granule cells are not the same cell type, but they can converge on related downstream assembly configurations.",
        "The central conclusion is developmental rather than taxonomic: dentate and cerebellar granule cells are not the same cell type, but distinct regional lineages can converge on related downstream postmitotic assembly configurations.",
    )
    body = body.replace(
        "The candidate-gene results support this interpretation.",
        "The candidate-gene results support this developmental interpretation.",
    )
    body = body.replace(
        "The sequencing data support the broad relevance of TGF-beta, BDNF, SMAD, MAPK, and related maturation pathways, but they revise the original conditioned-medium interpretation.",
        "The sequencing data support the broad relevance of TGF-beta, BDNF, SMAD, MAPK and related maturation pathways, but they convert the original conditioned-medium observation into a stage-window hypothesis.",
    )
    body = body.replace(
        "The sparse-coding model provides a mathematical reason why compact granule-like neurons may be useful.",
        "The sparse-coding model provides a mathematical reason why similar developmental assembly states may repeatedly resolve into compact granule-like designs.",
    )
    body = body.replace(
        "with a versioned archival DOI to be added here.",
        "with a versioned archival DOI assigned at release.",
    )
    body = body.replace(
        "### Limitations\n\n1. The study is computational and integrative.",
        "### Limitations\n\n1. The study is computational and integrative, and should be read as a developmental model with prioritized experimental tests.",
    )
    body = replace_section(body, "## Methods", "## References", METHODS_CLEAN.read_text(encoding="utf-8"))

    body = strip_pmids(body)
    body = normalize_gene_backticks(body)
    body = normalize_table_refs(body)

    declarations = """## Author Contributions

J.L. conceived the study question, curated and interpreted the biological framework, supervised the computational analyses, reviewed the outputs, prepared the figures and wrote the manuscript.

## Competing Interests

The author declares no competing or financial interests.

## Funding

No specific funding was declared.

"""
    body = body.replace("## References\n", declarations + "## References\n", 1)

    revised = normalize_table_refs(front + "\n" + intro + body)
    revised = re.sub(r"\n{3,}", "\n\n", revised).strip() + "\n"
    OUT_MD.write_text(revised, encoding="utf-8")
    return revised


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, name: str, size: float | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, before, after in [
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 10, 4),
        ("Heading 3", 12, 8, 2),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    if "Formula" not in [s.name for s in doc.styles]:
        formula = doc.styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = doc.styles["Formula"]
    formula.font.name = "Courier New"
    formula._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    formula.font.size = Pt(10)
    formula.paragraph_format.left_indent = Inches(0.25)
    formula.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    formula.paragraph_format.space_before = Pt(6)
    formula.paragraph_format.space_after = Pt(6)

    if "Reference" not in [s.name for s in doc.styles]:
        ref = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = doc.styles["Reference"]
    ref.font.name = "Times New Roman"
    ref._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    ref._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    ref.font.size = Pt(12)
    ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    ref.paragraph_format.space_after = Pt(6)
    ref.paragraph_format.first_line_indent = Inches(-0.25)
    ref.paragraph_format.left_indent = Inches(0.25)


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    ln = sect_pr.find(qn("w:lnNumType"))
    if ln is None:
        ln = OxmlElement("w:lnNumType")
        sect_pr.append(ln)
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")
    ln.set(qn("w:distance"), "360")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def parse_inline(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_font(run, "Times New Roman", 12)
        token = match.group(0)
        content = token
        bold = False
        italic = False
        font = "Times New Roman"
        if token.startswith("**"):
            content = token[2:-2]
            bold = True
        elif token.startswith("*"):
            content = token[1:-1]
            italic = True
        elif token.startswith("`"):
            content = token[1:-1]
            font = "Courier New"
        run = paragraph.add_run(content)
        set_font(run, font, 12 if font == "Times New Roman" else 10.5, bold=bold, italic=italic)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, "Times New Roman", 12)


def split_blocks(markdown: str) -> Iterable[str]:
    buffer: list[str] = []
    in_formula = False
    for line in markdown.splitlines():
        if line.strip() == "$$":
            if buffer and not in_formula:
                yield "\n".join(buffer).strip()
                buffer = []
            buffer.append(line)
            if in_formula:
                yield "\n".join(buffer).strip()
                buffer = []
                in_formula = False
            else:
                in_formula = True
            continue
        if in_formula:
            buffer.append(line)
            continue
        if not line.strip():
            if buffer:
                yield "\n".join(buffer).strip()
                buffer = []
            continue
        buffer.append(line.rstrip())
    if buffer:
        yield "\n".join(buffer).strip()


def add_markdown_to_docx(markdown: str, doc: Document) -> None:
    for block in split_blocks(markdown):
        if block.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(block[2:].strip())
            set_font(run, "Times New Roman", 16, bold=True)
        elif block.startswith("## "):
            p = doc.add_paragraph(block[3:].strip(), style="Heading 1")
        elif block.startswith("### "):
            p = doc.add_paragraph(block[4:].strip(), style="Heading 2")
        elif block.startswith("$$"):
            p = doc.add_paragraph(style="Formula")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(block)
            set_font(run, "Courier New", 10)
        elif all(line.startswith("- ") for line in block.splitlines()):
            for line in block.splitlines():
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                parse_inline(p, line[2:].strip())
        elif re.match(r"^\d+\. ", block):
            for line in block.splitlines():
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                parse_inline(p, re.sub(r"^\d+\. ", "", line).strip())
        elif block.startswith("Target journal:") or block.startswith("Date revised:"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            parse_inline(p, block.replace("  \n", "\n"))
        elif block.startswith("Figure ") or block.startswith("Supplementary Figure") or block.startswith("Graphical Abstract"):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            parse_inline(p, block)
        elif block.startswith("Ascoli") or re.match(r"^[A-Z][A-Za-z'\-]+, .+\(\d{4}\)\.", block):
            p = doc.add_paragraph(style="Reference")
            parse_inline(p, block)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            parse_inline(p, block.replace("\n", " "))


def build_reference_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    add_line_numbers(section)
    configure_styles(doc)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    doc.add_paragraph("Reference style template")
    doc.save(REFERENCE_DOCX)


def build_docx(markdown: str) -> None:
    build_reference_docx()
    subprocess.run(
        [
            "/opt/anaconda3/bin/pandoc",
            str(OUT_MD),
            "--from",
            "markdown+tex_math_dollars+tex_math_single_backslash",
            "--to",
            "docx",
            "--reference-doc",
            str(REFERENCE_DOCX),
            "--output",
            str(OUT_DOCX),
        ],
        check=True,
    )
    scrub_metadata_and_enable_line_numbers(OUT_DOCX)


def scrub_metadata_and_enable_line_numbers(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    core_name = "docProps/core.xml"
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == core_name:
                core = data.decode("utf-8")
                core = re.sub(r"<dc:creator>.*?</dc:creator>", "<dc:creator></dc:creator>", core)
                core = re.sub(r"<cp:lastModifiedBy>.*?</cp:lastModifiedBy>", "<cp:lastModifiedBy></cp:lastModifiedBy>", core)
                data = core.encode("utf-8")
            elif item.filename == "word/document.xml":
                document_xml = data.decode("utf-8")
                ln = '<w:lnNumType w:countBy="1" w:restart="continuous" w:distance="360"/>'
                if "<w:lnNumType" in document_xml:
                    document_xml = re.sub(r"<w:lnNumType\b[^>]*/>", ln, document_xml)
                else:
                    document_xml = re.sub(r"(<w:sectPr\b[^>]*>)", r"\1" + ln, document_xml)
                data = document_xml.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def main() -> int:
    markdown = build_submission_markdown()
    build_docx(markdown)
    print(f"Wrote Markdown: {OUT_MD}")
    print(f"Wrote DOCX: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
